import openpyxl
import datetime
import docx
import os
from docx.shared import RGBColor


def fill_table(units_list, document, department):

    table = document.add_table(rows=1, cols=6, style='Table Grid')
    hdr_cells = table.rows[0].cells  # returns tuple of cells in 0 row
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Наименование СИ'
    hdr_cells[2].text = 'Тип'
    hdr_cells[3].text = 'Зав. номер'
    hdr_cells[4].text = 'Годен до'
    hdr_cells[5].text = 'ФИО держателя'

    counter = 1
    for i in units_list:
        if i[6] == str(department):
            row_cells = table.add_row().cells
            row_cells[0].text = str(counter)
            row_cells[1].text = str(i[1])
            row_cells[2].text = str(i[2])
            row_cells[3].text = str(i[3])
            row_cells[4].text = str(i[4])
            row_cells[5].text = str(i[5])
            counter += 1

    first_col = table.columns[0].cells
    for i in first_col:
        i.width = 0.5 * 914400


def create_report(dicti, date):
    date_str_ = datetime.datetime.strftime(date, "%d.%m.%Y")
    for key in dicti:
        document = docx.Document()
        heading = document.add_heading().add_run('Служебная записка в отдел ' + key + ' от ' + date_str_)
        heading.font.color.rgb = RGBColor(0x0, 0x0, 0x0)
        document.add_paragraph('')

        if len(dicti[key][0]) != 0 and len(dicti[key][1]) != 0:
            document.add_paragraph('Настоящим сообщаю, что на сегодняшний день у Вас в отделе имеются в эксплуатации'
                                   ' средства измерения с истекшим сроком поверки:')
            fill_table(dicti[key][0], document, key)
            document.add_paragraph('\nКроме этого, в ближайшее время заканчивается поверка у следующих СИ Вашего отдела:')
            fill_table(dicti[key][1], document, key)
            document.add_paragraph('\nПрошу предоставить приборы на поверку.')

        elif len(dicti[key][0]) != 0 and len(dicti[key][1]) == 0:
            document.add_paragraph('Настоящим сообщаю, что на сегодняшний день у Вас в отделе имеются в эксплуатации'
                                   ' средства измерения с истекшим сроком поверки:')
            fill_table(dicti[key][0], document, key)
            document.add_paragraph('\nПрошу предоставить приборы на поверку.')

        elif len(dicti[key][0]) == 0 and len(dicti[key][1]) != 0:
            document.add_paragraph('Настоящим сообщаю, что в ближайшее время заканчивается поверка у следующих СИ Вашего отдела:')
            fill_table(dicti[key][1], document, key)
            document.add_paragraph('\nПрошу предоставить приборы на поверку.')

        else:
            document.add_paragraph('\nВсе приборы поверены!')

        # create a new folder for reports if it not exists
        if not os.path.isdir('reports_SI'):
            os.mkdir('reports_SI')

        date_str = datetime.datetime.strftime(date, "%Y.%m.%d")
        document.save('reports_SI\\' + date_str + ' сз в ' + str(key) + '.docx')


def main(file_path, days_left):

    if file_path == '':
        return 'File path is empty!'

    if type(file_path) != str:
        return 'String error!'

    wb = openpyxl.load_workbook(file_path)
    ws = wb.active

    data = tuple(ws.values)

    # for i in data:
    #     print(i)

    # if type(data[0][4]) == datetime.datetime:
    #     print('data[1][4]')

    days_left = days_left
    today = datetime.datetime.today()
    in_a_month = today + datetime.timedelta(days=days_left)
    expired = []
    soon_expire = []
    job_is_done = False
    for j in range(len(data[0])):
        # print(data[0][j])
        if data[0][j] == 'Годен до':
            for i in range(1, len(data)):
                if data[i][j] is None:
                    return 'Incorrect format of database: there is a blank in row: ' + str(i + 1)
                elif type(data[i][j]) != datetime.datetime:
                    return 'Incorrect format of database: values in column \'Годен до\' must be in datetime format'
                else:
                    if data[i][j] < today:
                        expired.append(list(data[i]))
                        expired[-1][j] = datetime.datetime.strftime(expired[-1][j], "%d.%m.%Y")
                    elif data[i][j] < in_a_month:
                        soon_expire.append(list(data[i]))
                        soon_expire[-1][j] = datetime.datetime.strftime(soon_expire[-1][j], "%d.%m.%Y")
            job_is_done = True
            break

    if not job_is_done:
        return 'Incorrect format of database: there is no title \'Годен до\''

    # print('expired:')
    # for i in expired:
    #     print(i)
    #
    # print('soon:')
    # for i in soon_expire:
    #     print(i)

    # find all departments which have expired devices
    departments_for_report = []
    for i in range(len(expired)):
        departments_for_report.append(str(expired[i][6]))

    for i in range(len(soon_expire)):
        departments_for_report.append(str(soon_expire[i][6]))

    departments_for_report = list(set(departments_for_report))  # remove duplicate values

    # format all units in dictionary in order: department: [[expired units][soon expired units]]
    dicti = {}
    for dep in departments_for_report:
        dep_list_exp = []
        dep_list_soon = []

        for unit in expired:
            if unit[6] == dep:
                dep_list_exp.append(unit)

        for unit in soon_expire:
            if unit[6] == dep:
                dep_list_soon.append(unit)

        dicti[dep] = [dep_list_exp, dep_list_soon]

    if len(dicti.keys()) != 0:
        create_report(dicti, today)
        return 'Report has been formed'
    else:
        return 'Everything is serviced!'


